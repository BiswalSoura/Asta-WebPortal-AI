from pathlib import Path

import pymupdf
from docx import Document

from app.knowledge.loaders import (
    get_document_loader,
)


def test_text_loader(
    tmp_path: Path,
) -> None:
    path = tmp_path / "sample.txt"

    path.write_text(
        "WebPortal help text",
        encoding="utf-8",
    )

    loader = get_document_loader(path)

    sections = loader.load(path)

    assert len(sections) == 1
    assert sections[0].content == (
        "WebPortal help text"
    )


def test_markdown_loader_preserves_heading(
    tmp_path: Path,
) -> None:
    path = tmp_path / "sample.md"

    path.write_text(
        "# Login\nEnter username and password.",
        encoding="utf-8",
    )

    loader = get_document_loader(path)

    sections = loader.load(path)

    assert len(sections) == 1
    assert sections[0].title == "Login"
    assert (
        "Enter username"
        in sections[0].content
    )


def test_docx_loader_preserves_heading(
    tmp_path: Path,
) -> None:
    path = tmp_path / "sample.docx"

    document = Document()

    document.add_heading(
        "Create New Project",
        level=1,
    )

    document.add_paragraph(
        "Enter the project owner title."
    )

    document.save(path)

    loader = get_document_loader(path)

    sections = loader.load(path)

    assert len(sections) == 1

    assert sections[0].title == (
        "Create New Project"
    )

    assert (
        "project owner title"
        in sections[0].content
    )


def test_pdf_loader_preserves_page_number(
    tmp_path: Path,
) -> None:
    path = tmp_path / "sample.pdf"

    document = pymupdf.open()

    page = document.new_page()

    page.insert_text(
        (72, 72),
        "WebPortal PDF content",
    )

    document.save(path)

    document.close()

    loader = get_document_loader(path)

    sections = loader.load(path)

    assert len(sections) == 1
    assert sections[0].page_number == 1
    assert (
        "WebPortal PDF content"
        in sections[0].content
    )